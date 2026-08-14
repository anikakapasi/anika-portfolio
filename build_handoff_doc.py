from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/anikakapasi/Desktop/website/Anika_Kapasi_Portfolio_Handoff.docx"

NAVY = "192A41"
BROWN = "763E2A"
TAN = "F1E7D3"
BLUE = "DCE8F2"
CREAM = "FCFBF7"
INK = "162238"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in kwargs[edge]:
                    element.set(qn("w:{}".format(key)), str(kwargs[edge][key]))


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Aptos", size=10.5, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_font(run, **kwargs)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(5)
    add_text(
        p,
        text,
        name="Georgia",
        size=16 if level == 1 else 12,
        color=BROWN,
        bold=True,
    )
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead and text.startswith(bold_lead):
        add_text(p, bold_lead, bold=True)
        add_text(p, text[len(bold_lead):])
    else:
        add_text(p, text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    add_text(p, text)
    return p


def add_rule(doc, color=BROWN):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_callout(doc, title, paragraphs, fill=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=130, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_border(cell, top={"val": "single", "sz": 8, "color": BROWN}, bottom={"val": "single", "sz": 8, "color": BROWN}, left={"val": "single", "sz": 8, "color": BROWN}, right={"val": "single", "sz": 8, "color": BROWN})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_text(p, title, name="Aptos", size=10.5, color=WHITE, bold=True)
    for item in paragraphs:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(2)
        add_text(p, "• ", color=WHITE)
        add_text(p, item, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_two_column_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.72)
    table.columns[1].width = Inches(5.75)
    set_repeat_table_header(table.rows[0])
    header = table.rows[0].cells
    for cell, text in zip(header, ("File / area", "Current state and continuation notes")):
        shade(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        add_text(p, text, color=WHITE, bold=True, size=9.5)
    for area, detail in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell, top=90, bottom=90)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell_border(cell, bottom={"val": "single", "sz": 4, "color": "D8D1C8"})
        shade(cells[0], TAN)
        p = cells[0].paragraphs[0]
        add_text(p, area, name="Georgia", size=10, color=BROWN, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.line_spacing = 1.1
        add_text(p, detail, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_link_paragraph(doc, label, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_text(p, label + ": ", bold=True)
    add_text(p, url, color=BROWN, italic=True)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.76)
    section.right_margin = Inches(0.76)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(hp, "ANIKA KAPASI  /  PORTFOLIO HANDOFF", size=8.5, color=BROWN, bold=True)

    # Footer with page numbering.
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(fp, "Portfolio continuation handoff  •  ", size=8, color=BROWN)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    add_text(title, "Anika Kapasi Portfolio", name="Georgia", size=28, color=BROWN, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(7)
    add_text(subtitle, "Continuation handoff • August 13, 2026", size=11, color=NAVY, italic=True)
    add_rule(doc)

    add_body(doc, "Purpose: Use this document to continue the existing personal journalism portfolio in a new task without losing the site’s visual direction, content decisions, and high-priority open work.", bold_lead="Purpose: ")
    add_body(doc, "Project folder: /Users/anikakapasi/Desktop/website")
    add_body(doc, "Primary pages: index.html (Home), selected-works.html (Selected Works), data-projects.html (Data-Driven Projects), and congress-data-analysis.html (Congress project detail).")

    add_callout(doc, "START HERE — NON-NEGOTIABLE CURRENT REQUEST", [
        "Remove every decorative blue, brown, or black rectangle/underlay that sits behind article cards on Selected Works or project cards on Data-Driven Projects.",
        "Do not add those under-card backdrops again. Header and carousel framing may remain; card grids must sit cleanly on the page.",
        "Preserve existing content and local photo assets. Inspect the current source before editing and verify desktop plus responsive layouts afterward.",
    ])

    add_heading(doc, "Visual system to preserve")
    for item in [
        "Base palette: warm off-white pages, deep navy About/footer area, dark brown display type and rules, pale blue geometry, and light tan card surfaces.",
        "Typography: editorial serif for display and story text; fine script only for the About Me title; restrained uppercase sans-serif for metadata and labels.",
        "Use the geometric theme intentionally: thin brown outlines, brown top rules, pale blue/tan framing in headers or carousel. Keep it subtle; do not use it as a backdrop under cards.",
        "Navigation convention: AK at left; Home, Selected Works, Data Projects at right. External/article arrows use a diagonal up-right arrow (↗).",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Page-by-page status")
    add_two_column_table(doc, [
        ("index.html\nHome + About", "Home is finalized around the option 8 / mirror-collage direction. The About section is deep navy with the scripted About Me title, biography at left and photo carousel at right. Preserve the current hierarchy and photo assets; carousel transitions should feel like a smooth horizontal swipe rather than rigid/fade changes."),
        ("selected-works.html\nSelected Works", "Title is SELECTED WORKS. The intro label/dek were removed. Filters are a single-selection system: clicking a new filter clears the previous one. Cards are newest-first, tan, with brown top rules, dark-brown editorial titles, italic dates above dividers and consistent publication lines. Remove all card underlays/backdrops now."),
        ("data-projects.html\nData-Driven Projects", "Header title stays on one line. Its pale-blue and brown-outline decoration should be straight, not tilted. Project cards are a two-column grid, newest first, standardized so every footer stays inside the tan card. Remove all under-card blue/brown/black decorative rectangles."),
        ("congress-data-analysis.html\nDetail page", "Simplified title page with no author name, no “Data Analysis – Python” subtitle, and no decorative title boxes. Body/list type is 18px, left aligned. Remaining layout work focuses on equal narrative spacing, course-note alignment, and a small clear tan gap between the two graph frames."),
    ])

    add_heading(doc, "Home / About: approved content and behavior")
    add_body(doc, "Bio copy (current):")
    for text in [
        "I’m a journalism and data science major at Boston University, originally from the Bay Area, California. I’m passionate about data-driven stories that encourage transparency and accountability in the communities I report on. I’ve been writing and freelancing since high school while continuing to build the technical skills in my classes to better analyze and understand data.",
        "My experience at the California Scholastic Press Association workshop in 2024, which reinforced news fundamentals under the rush of tight deadlines, made me realize that journalism was the career I wanted to pursue seriously. As the industry evolves in an increasingly AI-driven landscape, the news itself isn’t going anywhere … just the form in which we deliver it will. I’m excited to be part of shaping what comes next for the profession.",
        "Outside of journalism, you’ll probably find me listening to live jazz or classical music, taking photos, or convincing myself that going on a run is a fun way to spend my free time.",
        "I’m always looking for opportunities to report for and grow with local publications in the Boston area. Feel free to reach out at anika.kapasi@gmail.com!",
    ]:
        add_bullet(doc, text)
    add_body(doc, "About metadata:")
    for item in [
        "Interests: Technology · Policy · Finance · Data Journalism",
        "Words In: The Daily Free Press · Los Gatan · Los Altos Town Crier · The Outlook",
        "Currently: Co-City Editor · The Daily Free Press; Teaching Assistant · Computational Journalism, Prof. Brooke Williams; Freelancer · Los Gatan",
        "Use middle-dot separators (·), not em dashes. Text should remain left-aligned in its intended columns.",
    ]:
        add_bullet(doc, item)
    add_body(doc, "Carousel notes:")
    for item in [
        "Use local assets in the website, not absolute Downloads paths. Keep object-fit/cropping intentional and always fill the frame without letterbox/black bars.",
        "The newspaper slide must fill the photo frame with no black edges; emphasize the front page/top portion.",
        "The carousel should animate horizontally with eased sliding and support a swipe-like interaction. It should not feel like a rigid replacement.",
        "Current caption mapping: 04 First time covering a breaking news event; 05 Meeting Boston Mayor Michelle Wu at the first State of the Schools Address; 06 Selfie with my first official press pass!; 07 Last publication as Editor-in-Chief of my high school paper — Senior Magazine edition; 08 Spending two weeks at the California Scholastic Press Association workshop learning news fundamentals and strong journalistic practices; 09 Presenting two articles at a student showcase held by the Society of Professional Journalists, BU chapter; 11 Editing layouts during print night as a co-Editor-in-Chief; 12 One of my first print nights as a new staffer on my high school paper; 13 First front page story for the FreeP.",
        "Photos 01–03 and photo 10 were removed. A counselor photo was inserted after slide 06 with the caption: Returning as a counselor to the California Scholastic Press Association workshop.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Selected Works: content and behavior")
    for item in [
        "Filters shown: ALL, FAVORITES, BREAKING NEWS, ENTERPRISE, FEATURES, SOLUTIONS, POLITICS & GOVERNMENT, BUSINESS & ECONOMY, TECHNOLOGY, EDUCATION, COMMUNITY & CULTURE. Only one filter can be active at once.",
        "Article cards must remain photo-forward and consistent; do not introduce decorative slabs beneath them.",
        "Add / retain this affordable-housing article with Business & Economy, Community & Culture, and Politics & Government tags: https://www.losaltosonline.com/business/los-altos-first-fully-affordable-housing-development-takes-aim-at-carbon-emissions/article_a3aa434c-9cb4-4800-85be-d280a4275fa5.html. Its photo is in the portfolio photo folder.",
        "The Viva la Musica card was removed. Where a publication was shown as OUTLOOK, use The Outlook.",
    ]:
        add_bullet(doc, item)
    add_body(doc, "Selected Works footer / All Bylines:")
    add_link_paragraph(doc, "The Daily Free Press", "https://dailyfreepress.com/staff_name/anika-kapasi/")
    add_link_paragraph(doc, "Los Gatan", "https://losgatan.com/guest-author/anika+kapasi/")
    add_link_paragraph(doc, "Los Altos Town Crier", "https://www.losaltosonline.com/search/?l=25&sd=desc&s=start_time&f=html&t=article%2Cvideo%2Cyoutube%2Ccollection&q=anika+kapasi")
    add_body(doc, "Footer label is All Bylines at approximately 20px. Publication names should retain their own links and be separated using middle dots.")

    add_heading(doc, "Data-Driven Projects: required structure")
    for item in [
        "Two cards per row, ordered most recent to least recent; cards must use matching overall dimensions while holding all their information within the tan surface.",
        "Card sequence: tags → title (32px) → description (15px; use “Description goes here” where missing) → date immediately above divider → footer below divider. No element may fall outside its card.",
        "Santa Clara County continues aging — date July 15, 2026; tags PYTHON · PANDAS · DEMOGRAPHICS; source Los Gatan; description remains placeholder for now.",
        "A Look at Age, Tenure, and Power in the 118th Congress — date Fall 2025; tags PYTHON · PANDAS · MACHINE LEARNING; description placeholder for now; footer exactly: COMPLETED AS PART OF BU DS110 (INTRODUCTION TO DS WITH PYTHON). Keep that footer one line and the same small size as LOS GATAN.",
        "Header rectangles (blue and brown outline) should be square/straight with no rotation. Pale-blue header rectangle was recently moved upward.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Congress data-analysis detail page: remaining precision work")
    for item in [
        "Hero title: A Look at Age, Tenure, and Power in the 118th Congress. Target size 60px. No author byline, no title decoration, and no “(Data Analysis – Python)”.",
        "Purpose, My Work heading, and list are 18px and left aligned. The course note must start exactly at the same left edge as narrative paragraphs (the user repeatedly flagged this).",
        "Course note text: Completed as part of Boston University DS110 (Introduction to Data Science with Python).",
        "Make the vertical gaps between all findings narrative paragraphs consistent. The first Findings paragraph through the conclusion and the following Furthermore paragraph should read as a continuous column with the same gap between each paragraph.",
        "Move the second graph / Distribution of Age (118th Congress) downward independently enough to create visible tan space between its frame and the first graph’s blue/brown framing. It should not be forced by the left copy’s height.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Working conventions for the next task")
    for item in [
        "Inspect the actual code first; this handoff describes the design intent and high-priority decisions, but the current source is authoritative for exact placement and asset filenames.",
        "Use apply_patch for source edits. Do not reset or delete user changes unrelated to the specific request.",
        "Use responsive CSS rather than one-off fixed viewport hacks. Check wide desktop and a narrower/mobile layout after visual edits.",
        "When fitting images, prefer object-fit: cover with an explicit object-position per photo. Never use black letterboxing unless expressly requested.",
    ]:
        add_bullet(doc, item)

    add_callout(doc, "SUGGESTED FIRST MESSAGE IN THE NEW TASK", [
        "Continue in /Users/anikakapasi/Desktop/website. Read the handoff document first and inspect the current source before editing.",
        "Highest priority: remove every decorative blue, brown, or black rectangle/underlay behind cards from selected-works.html and data-projects.html, and do not reintroduce them.",
        "Then verify the two card grids, header decorations, and responsive behavior without changing completed content or photo assets unnecessarily.",
    ], fill=BROWN)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
